import streamlit as st
import sqlite3
import pandas as pd
from datetime import datetime, timedelta
from streamlit_calendar import calendar
import os
import shutil
import time
import json
import re
import uuid
import wave
import struct
import gzip
import hashlib
import secrets as py_secrets
import hmac
import smtplib
import ssl
import zipfile
import base64
import tempfile
from io import BytesIO
from email.message import EmailMessage
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from openai import OpenAI

try:
    import audioop  # removed in newer Python versions; optional in this app
except Exception:
    audioop = None

def _load_dotenv(dotenv_path: str = ".env") -> None:
    """Load simple KEY=VALUE pairs from .env into os.environ (no external deps)."""
    if not os.path.exists(dotenv_path):
        return
    try:
        with open(dotenv_path, "r", encoding="utf-8") as fh:
            for raw_line in fh:
                line = raw_line.strip()
                if not line or line.startswith("#"):
                    continue
                if "=" not in line:
                    continue
                key, value = line.split("=", 1)
                key = key.strip()
                value = value.strip().strip("'").strip('"')
                if not key:
                    continue
                os.environ.setdefault(key, value)
    except Exception:
        return


def _get_setting(name: str, default: str = "") -> str:
    """Read settings from Streamlit secrets first, then env vars."""
    try:
        if hasattr(st, "secrets") and name in st.secrets:
            return str(st.secrets[name])
    except Exception:
        pass
    return os.getenv(name, default)


_load_dotenv()

try:
    import pypandoc
except Exception:
    pypandoc = None

try:
    import win32com.client
    import pythoncom
except Exception:
    win32com = None
    pythoncom = None

HAS_PYPANDOC = pypandoc is not None
HAS_WIN32_COM = win32com is not None

# --- 配置区 ---
LEGACY_UPLOAD_DIR = "uploads"
LEGACY_BACKUP_DIR = "backups"
LEGACY_DB_PATH = "my_lab_data.db"
DATA_DIR = "data"


def _get_streamlit_user_email() -> str | None:
    """Best-effort: return the email of the signed-in Streamlit user (Cloud/private apps)."""
    try:
        if hasattr(st, "user"):
            info = st.user.to_dict() or {}
            email = info.get("email")
            if email:
                return str(email).strip().lower()
    except Exception:
        pass
    return None


def _parse_csv_list(value: str) -> list[str]:
    if not value:
        return []
    parts = re.split(r"[,\s]+", str(value))
    return [p.strip().lower() for p in parts if p.strip()]


def _get_auth_settings() -> dict:
    mode = str(_get_setting("LAB_DIARY_AUTH_MODE", "")).strip().lower()  # "email_otp" | "password" | ""
    allowed_domains = _parse_csv_list(_get_setting("LAB_DIARY_ALLOWED_EMAIL_DOMAINS", ""))
    allowed_emails = _parse_csv_list(_get_setting("LAB_DIARY_ALLOWED_EMAILS", ""))
    session_minutes = int(str(_get_setting("LAB_DIARY_AUTH_SESSION_MINUTES", "720")).strip() or "720")
    code_minutes = int(str(_get_setting("LAB_DIARY_AUTH_CODE_MINUTES", "10")).strip() or "10")
    dev_show_code = str(_get_setting("LAB_DIARY_AUTH_DEV_SHOW_CODE", "0")).strip().lower() in ("1", "true", "yes")
    shared_password = str(_get_setting("LAB_DIARY_SHARED_PASSWORD", "")).strip()
    shared_password_hash = str(_get_setting("LAB_DIARY_SHARED_PASSWORD_HASH", "")).strip().lower()

    smtp_host = str(_get_setting("SMTP_HOST", "")).strip()
    smtp_port = int(str(_get_setting("SMTP_PORT", "587")).strip() or "587")
    smtp_user = str(_get_setting("SMTP_USER", "")).strip()
    smtp_password = str(_get_setting("SMTP_PASSWORD", "")).strip()
    smtp_from = str(_get_setting("SMTP_FROM", "")).strip()
    smtp_use_tls = str(_get_setting("SMTP_USE_TLS", "1")).strip().lower() in ("1", "true", "yes")
    smtp_use_ssl = str(_get_setting("SMTP_USE_SSL", "0")).strip().lower() in ("1", "true", "yes")
    smtp_timeout = int(str(_get_setting("SMTP_TIMEOUT", "20")).strip() or "20")
    smtp_debug = str(_get_setting("SMTP_DEBUG", "0")).strip().lower() in ("1", "true", "yes")

    return {
        "mode": mode,
        "allowed_domains": allowed_domains,
        "allowed_emails": allowed_emails,
        "session_minutes": session_minutes,
        "code_minutes": code_minutes,
        "dev_show_code": dev_show_code,
        "shared_password": shared_password,
        "shared_password_hash": shared_password_hash,
        "smtp_host": smtp_host,
        "smtp_port": smtp_port,
        "smtp_user": smtp_user,
        "smtp_password": smtp_password,
        "smtp_from": smtp_from,
        "smtp_use_tls": smtp_use_tls,
        "smtp_use_ssl": smtp_use_ssl,
        "smtp_timeout": smtp_timeout,
        "smtp_debug": smtp_debug,
    }


def _is_allowed_login_email(email: str, auth: dict) -> bool:
    email = (email or "").strip().lower()
    if not email or "@" not in email:
        return False
    if auth["allowed_emails"] and email not in auth["allowed_emails"]:
        return False
    if auth["allowed_domains"]:
        domain = email.split("@", 1)[1]
        if domain not in auth["allowed_domains"]:
            return False
    return True


def _send_email_login_code(email: str, code: str, auth: dict) -> None:
    if auth["dev_show_code"]:
        return
    if not auth["smtp_host"] or not auth["smtp_from"]:
        raise RuntimeError("Missing SMTP configuration (SMTP_HOST/SMTP_FROM).")

    msg = EmailMessage()
    msg["Subject"] = "Lab Diary AI 登录验证码"
    msg["From"] = auth["smtp_from"]
    msg["To"] = email
    msg.set_content(
        f"你的登录验证码是：{code}\n\n"
        f"有效期：{auth['code_minutes']} 分钟\n\n"
        "如非本人操作，请忽略此邮件。"
    )

    host = auth["smtp_host"]
    port = auth["smtp_port"]
    timeout = auth["smtp_timeout"]
    context = ssl.create_default_context()

    step = "connect"
    try:
        if auth["smtp_use_ssl"]:
            server = smtplib.SMTP_SSL(host, port, timeout=timeout, context=context)
        else:
            server = smtplib.SMTP(host, port, timeout=timeout)
        try:
            if auth["smtp_debug"]:
                server.set_debuglevel(1)
            step = "ehlo"
            server.ehlo()
            if auth["smtp_use_tls"] and not auth["smtp_use_ssl"]:
                step = "starttls"
                server.starttls(context=context)
                step = "ehlo_after_starttls"
                server.ehlo()
            if auth["smtp_user"] and auth["smtp_password"]:
                step = "login"
                server.login(auth["smtp_user"], auth["smtp_password"])
            step = "send_message"
            server.send_message(msg)
        finally:
            try:
                step = "quit"
                server.quit()
            except Exception:
                pass
    except Exception as exc:
        raise RuntimeError(
            f"SMTP failed at step={step}. "
            f"Check SMTP_HOST/PORT, TLS vs SSL (587=STARTTLS, 465=SSL), and whether the hosting platform blocks outbound SMTP. "
            f"Original error: {exc!r}"
        ) from exc


def _clear_auth_session() -> None:
    for key in list(st.session_state.keys()):
        if key.startswith("auth_"):
            del st.session_state[key]


def require_app_login() -> None:
    """
    In-app login page (email OTP).
    Enable by setting `LAB_DIARY_AUTH_MODE=email_otp` in Streamlit Cloud secrets.
    """
    auth = _get_auth_settings()
    if auth["mode"] not in ("email_otp", "password"):
        return

    now = datetime.now()
    authed_email = (st.session_state.get("auth_email") or "").strip().lower()
    expires_at = st.session_state.get("auth_expires_at")
    if authed_email and isinstance(expires_at, datetime) and expires_at > now:
        return

    st.set_page_config(page_title="Lab Diary AI 登录", layout="centered", page_icon="🔐")
    st.title("🔐 登录 Lab Diary AI")
    if auth["mode"] == "email_otp":
        st.caption("输入邮箱获取验证码登录。")
    else:
        st.caption("输入邮箱 + 共享口令登录（无需发验证码）。")

    email = st.text_input("邮箱", key="auth_input_email", placeholder="name@domain.com").strip().lower()
    if email and not _is_allowed_login_email(email, auth):
        if auth["allowed_domains"] or auth["allowed_emails"]:
            st.error("该邮箱不在允许范围内。")
        else:
            st.warning("当前未配置允许邮箱/域名白名单：任何邮箱都可以请求验证码。建议设置 `LAB_DIARY_ALLOWED_EMAIL_DOMAINS`。")

    if auth["mode"] == "email_otp":
        col1, col2 = st.columns(2)
        with col1:
            if st.button("发送验证码", use_container_width=True):
                if not email or "@" not in email:
                    st.error("请输入正确的邮箱。")
                elif not _is_allowed_login_email(email, auth):
                    st.error("该邮箱不在允许范围内。")
                else:
                    code = f"{py_secrets.randbelow(1000000):06d}"
                    st.session_state["auth_pending_email"] = email
                    st.session_state["auth_code_hash"] = hashlib.sha256(code.encode("utf-8")).hexdigest()
                    st.session_state["auth_code_expires_at"] = now + timedelta(minutes=auth["code_minutes"])
                    st.session_state["auth_code_sent_at"] = now
                    try:
                        _send_email_login_code(email, code, auth)
                        if auth["dev_show_code"]:
                            st.info(f"DEV 模式：验证码是 {code}")
                        st.success("验证码已发送，请查收邮箱。")
                    except Exception as exc:
                        st.error(f"发送失败：{exc}")

        with col2:
            code_input = st.text_input("验证码", key="auth_input_code", placeholder="6 位数字").strip()
            if st.button("登录", type="primary", use_container_width=True):
                pending_email = (st.session_state.get("auth_pending_email") or "").strip().lower()
                code_hash = st.session_state.get("auth_code_hash")
                code_expires = st.session_state.get("auth_code_expires_at")

                if not pending_email:
                    st.error("请先发送验证码。")
                elif not isinstance(code_expires, datetime) or code_expires <= now:
                    st.error("验证码已过期，请重新发送。")
                elif not code_input or not code_input.isdigit() or len(code_input) != 6:
                    st.error("请输入 6 位数字验证码。")
                else:
                    input_hash = hashlib.sha256(code_input.encode("utf-8")).hexdigest()
                    if input_hash != code_hash:
                        st.error("验证码错误。")
                    else:
                        st.session_state["auth_email"] = pending_email
                        st.session_state["auth_expires_at"] = now + timedelta(minutes=auth["session_minutes"])
                        for k in ("auth_pending_email", "auth_code_hash", "auth_code_expires_at", "auth_code_sent_at", "auth_input_code"):
                            if k in st.session_state:
                                del st.session_state[k]
                        st.success("登录成功。")
                        st.rerun()
    else:
        password_input = st.text_input("共享口令", key="auth_input_password", type="password").strip()
        if st.button("登录", type="primary", use_container_width=True):
            if not email or "@" not in email:
                st.error("请输入正确的邮箱。")
            elif not _is_allowed_login_email(email, auth):
                st.error("该邮箱不在允许范围内。")
            elif not password_input:
                st.error("请输入共享口令。")
            elif not auth.get("shared_password") and not auth.get("shared_password_hash"):
                st.error("未配置共享口令：请在 Secrets 中设置 `LAB_DIARY_SHARED_PASSWORD` 或 `LAB_DIARY_SHARED_PASSWORD_HASH`。")
            else:
                ok = False
                if auth.get("shared_password_hash"):
                    digest = hashlib.sha256(password_input.encode("utf-8")).hexdigest().lower()
                    ok = hmac.compare_digest(digest, auth["shared_password_hash"])
                elif auth.get("shared_password"):
                    ok = hmac.compare_digest(password_input, auth["shared_password"])
                if not ok:
                    st.error("共享口令错误。")
                else:
                    st.session_state["auth_email"] = email
                    st.session_state["auth_expires_at"] = now + timedelta(minutes=auth["session_minutes"])
                    for k in ("auth_pending_email", "auth_code_hash", "auth_code_expires_at", "auth_code_sent_at", "auth_input_code", "auth_input_password"):
                        if k in st.session_state:
                            del st.session_state[k]
                    st.success("登录成功。")
                    st.rerun()

    st.divider()
    if auth["mode"] == "email_otp":
        st.caption("需要帮助？请联系管理员配置邮件服务器（SMTP）和允许域名白名单。")
    else:
        st.caption("需要帮助？请联系管理员配置共享口令和允许域名白名单。")
    st.stop()


def get_storage_paths() -> dict:
    """
    Multi-user isolation:
    - If Streamlit provides a signed-in user email, store data in `data/users/<hash>/`.
    - Otherwise (local/dev), fall back to legacy paths in the repo root.
    """
    session_email = str(st.session_state.get("auth_email", "")).strip().lower()
    override_email = _get_setting("LAB_DIARY_USER_EMAIL", "").strip().lower()
    email = session_email or override_email or _get_streamlit_user_email()
    if email:
        digest = hashlib.sha256(email.encode("utf-8")).hexdigest()[:16]
        root = os.path.join(DATA_DIR, "users", digest)
        upload_dir = os.path.join(root, "uploads")
        backup_dir = os.path.join(root, "backups")
        db_path = os.path.join(root, "my_lab_data.db")
        user_label = email
    else:
        root = "."
        upload_dir = LEGACY_UPLOAD_DIR
        backup_dir = LEGACY_BACKUP_DIR
        db_path = LEGACY_DB_PATH
        user_label = "local"
    os.makedirs(upload_dir, exist_ok=True)
    os.makedirs(backup_dir, exist_ok=True)
    return {
        "user_label": user_label,
        "root": root,
        "upload_dir": upload_dir,
        "backup_dir": backup_dir,
        "db_path": db_path,
    }

# --- 语音识别（暂时下线）---
# 你之前配置的火山引擎语音识别相关代码已单独存档，方便之后恢复：
# 见 `archived/volc_asr_reference.py`

# --- DeepSeek AI 配置 ---
DEEPSEEK_API_KEY = _get_setting("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = _get_setting("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEEPSEEK_MODEL = _get_setting("DEEPSEEK_MODEL", "deepseek-chat")

# --- 设计风格配置 ---
# 色彩系统
COLORS = {
    'primary': '#1e293b',      # 深蓝灰 - 主色
    'secondary': '#334155',    # 中蓝灰 - 辅助色
    'background': '#f8fafc',   # 极浅灰 - 背景
    'accent': '#3b82f6',       # 科技蓝 - 强调
    'success': '#10b981',      # 成功绿
    'warning': '#f59e0b',      # 警告橙
    'error': '#ef4444',        # 错误红
    'info': '#06b6d4',         # 信息蓝
    'research': '#8b5cf6',     # 科研紫
    'clinical': '#06b6d4',     # 临床青
    'course': '#f59e0b',       # 课程金
    'other': '#6b7280',        # 中性灰
}

# 字体
FONTS = {
    'family': 'Inter, -apple-system, BlinkMacSystemFont, sans-serif',
    'mono': 'JetBrains Mono, monospace'
}

for d in [LEGACY_UPLOAD_DIR, LEGACY_BACKUP_DIR]:
    os.makedirs(d, exist_ok=True)

# ==================== 工具函数 ====================
def get_versioned_upload_path(filename):
    """Return upload path plus versioned filename to avoid overwriting."""
    upload_dir = get_storage_paths()["upload_dir"]
    base, ext = os.path.splitext(filename)
    candidate = filename
    idx = 1
    while os.path.exists(os.path.join(upload_dir, candidate)):
        candidate = f"{base}_v{idx}{ext}"
        idx += 1
    return os.path.join(upload_dir, candidate), candidate

def normalize_task_row(row):
    if isinstance(row, dict):
        return row
    if hasattr(row, "to_dict"):
        return row.to_dict()
    return row

def sanitize_filename(value, default="record"):
    value = re.sub(r"[^\w\-]+", "_", value).strip("_")
    return value or default

def shorten_task_name(name: str, max_length: int = 28) -> str:
    """强制任务名简洁且去除冗余标点"""
    if not name:
        return "未命名任务"
    clean = re.sub(r"\s+", " ", str(name)).strip()
    clean = re.sub(r"[，。,.；;、：:]+$", "", clean)
    if len(clean) > max_length:
        clean = clean[:max_length].rstrip() + "…"
    return clean

# ==================== AI 功能 ====================
def get_ai_client():
    """获取 DeepSeek AI 客户端"""
    if not DEEPSEEK_API_KEY:
        return None
    return OpenAI(base_url=DEEPSEEK_BASE_URL, api_key=DEEPSEEK_API_KEY)

def ai_polish_text(client, text, extra_instruction=None):
    """AI 润色功能"""
    if not text:
        return "请先输入文本。"
    user_content = text
    if extra_instruction:
        user_content = f"{text}\n\n[补充要求]\n{extra_instruction}"
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": "你是一个神经科学助手。请将用户的实验记录润色为学术风格。直接输出结果。"},
                {"role": "user", "content": user_content}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {e}"

def ai_extract_metadata(client, text):
    """
    AI只提取元数据，不修改原始内容
    返回: {date, task_name, category, tags}
    """
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": """你是一个元数据提取助手。从实验记录中提取以下信息，但不要修改原始记录内容：
1. 日期（YYYY-MM-DD格式）
2. 任务/实验名称
3. 类别（科研/临床/课程/其他）
4. 标签（以#开头，多个标签用空格分隔）

只返回JSON格式，不要添加解释。"""},
                {"role": "user", "content": text[:2000]}  # 限制输入长度
            ],
            response_format={"type": "json_object"}
        )
        content = response.choices[0].message.content
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]
        return json.loads(content)
    except Exception as e:
        print(f"AI metadata extraction error: {e}")
        return {}

def ai_parse_schedule(client, text, attachment_notes=None):
    """将大白话转换为结构化的 JSON 任务列表"""
    today_str = datetime.now().strftime("%Y-%m-%d")
    weekday_str = datetime.now().strftime("%A")
    
    system_prompt = f"""
    You are a smart scheduler. Today is {today_str} ({weekday_str}).
    Extract tasks from the user's natural language description.
    
    Rules:
    1. Calculate exact dates (e.g., "next Friday", "tomorrow", "for 3 days").
    2. Return a JSON object containing a list under key "tasks".
    3. Each task must have: "date" (YYYY-MM-DD), "task_name", "category" (choose from: 科研, 临床, 课程, 其他), "tags" (string starting with #).
    4. Include "record_outline": 1-3 sentences summarizing what should be recorded in the experiment log (关键材料/操作/观察) for this task.
    5. Do not output markdown code blocks, just raw JSON.
    """
    
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"{text}\n\n[附件参考]\n{attachment_notes}" if attachment_notes else text}
            ],
            response_format={"type": "json_object"}
        )
        content = response.choices[0].message.content
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]
        data = json.loads(content)
        return data.get("tasks", [])
    except Exception as e:
        print(f"AI Parse Error: {e}")
        return []

def ai_generate_weekly_report(client, records, start_date, end_date):
    """基于近 7 天的记录自动生成周报内容"""
    if not records:
        return ""
    bullet_lines = []
    for row in records:
        snippet = (row.get("details") or "").replace("\n", " ")
        snippet = re.sub(r"\s+", " ", snippet)
        if len(snippet) > 200:
            snippet = snippet[:200] + "…"
        bullet_lines.append(f"- {row.get('date', '')} {row.get('task_name', '')}：{snippet}")
    prompt = f"""你是科研助理，请将以下 {len(records)} 条实验记录整理为一篇结构化的科研周报，突出关键进展、问题与下一步计划。
时间区间：{start_date} ~ {end_date}

{chr(10).join(bullet_lines)}
"""
    try:
        resp = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": "你是经验丰富的实验室PI，擅长将记录整理成周报，语言简洁专业。"},
                {"role": "user", "content": prompt}
            ]
        )
        return resp.choices[0].message.content
    except Exception as exc:
        return f"生成失败：{exc}"

# ==================== 文档处理 ====================
FONT_CANDIDATES = [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\msyh.ttf",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simfang.ttf",
]
TEXT_LIKE_EXTS = {".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".yaml", ".yml", ".log"}
LEGACY_TEXT_EXTS = {".md", ".markdown", ".txt", ".docx", ".doc", ".rtf"}
LEGACY_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".svg"}
IMAGE_MIME_MAP = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".svg": "image/svg+xml",
}

def _decode_text_preview(data: bytes, max_chars: int = 1500) -> str:
    """尝试多种编码获取文本片段"""
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="ignore")
    text = text.strip()
    return (text[:max_chars] + "…") if len(text) > max_chars else text

def _decode_text_full(data: bytes, strip: bool = True) -> str:
    """尝试多种编码解码为完整字符串"""
    for enc in ("utf-8", "utf-16", "gbk", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="ignore")
    return text.strip() if strip else text

def _persist_image_as_markdown(data: bytes, original_name: str) -> str:
    """保存图片到 uploads，并返回内联 Markdown"""
    ext = os.path.splitext(original_name)[1].lower()
    if ext not in IMAGE_MIME_MAP:
        ext = ".png"
    base = sanitize_filename(os.path.splitext(original_name)[0] or "legacy_image")
    filename = f"{base}{ext}"
    save_path, stored_name = get_versioned_upload_path(filename)
    with open(save_path, "wb") as fh:
        fh.write(data)
    mime = IMAGE_MIME_MAP.get(ext, "application/octet-stream")
    payload = base64.b64encode(data).decode("ascii")
    data_uri = f"data:{mime};base64,{payload}"
    note_path = save_path.replace("\\", "/")
    return f"![{stored_name}]({data_uri})\n\n_原图已保存：{note_path}_"

def docx_to_markdown_with_assets(docx_bytes: bytes, origin_name: str) -> str:
    """将 DOCX 转 Markdown，保留段落、表格、图片"""
    doc = Document(BytesIO(docx_bytes))
    lines = []
    for block in _iter_docx_block_items(doc):
        if isinstance(block, Paragraph):
            chunk = _docx_paragraph_to_markdown(block)
            if chunk:
                lines.append(chunk)
        elif isinstance(block, Table):
            table_md = _docx_table_to_markdown(block)
            if table_md:
                lines.append(table_md)
    image_lines = _collect_docx_image_markdown(docx_bytes, origin_name)
    if image_lines:
        lines.append("### 附件图片")
        lines.extend(image_lines)
    return "\n\n".join(lines).strip()

def convert_document_bytes_to_markdown(data_bytes: bytes, origin_name: str, ext: str) -> str:
    """统一入口：将 doc/docx/rtf 转为 Markdown"""
    ext = ext.lower()
    if ext == ".docx":
        return docx_to_markdown_with_assets(data_bytes, origin_name)
    if ext == ".doc":
        converted = _pandoc_convert(data_bytes, ".doc", "docx")
        if not converted:
            converted = _convert_doc_via_win32(data_bytes)
        if converted:
            return docx_to_markdown_with_assets(converted, origin_name)
        fallback = _pandoc_convert(data_bytes, ".doc", "md")
        if fallback:
            return fallback.decode("utf-8")
        return ""
    if ext == ".rtf":
        converted = _pandoc_convert(data_bytes, ".rtf", "md")
        if converted:
            return converted.decode("utf-8")
        return _decode_text_full(data_bytes)
    return ""

# ==================== 数据库操作 ====================
def get_db_connection():
    db_path = get_storage_paths()["db_path"]
    conn = sqlite3.connect(db_path, timeout=30)
    try:
        conn.execute("PRAGMA journal_mode=WAL;")
    except Exception:
        pass
    return conn

def init_and_migrate_db():
    conn = get_db_connection()
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            date TEXT, task_name TEXT, category TEXT, is_done INTEGER, details TEXT, tags TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # 添加字段（如果不存在）
    c.execute("PRAGMA table_info(tasks)")
    cols = [i[1] for i in c.fetchall()]
    if 'tags' not in cols:
        try:
            c.execute("ALTER TABLE tasks ADD COLUMN tags TEXT DEFAULT ''")
            conn.commit()
        except:
            pass
    conn.commit()
    conn.close()

def auto_backup():
    paths = get_storage_paths()
    db_path = paths["db_path"]
    backup_dir = paths["backup_dir"]
    if os.path.exists(db_path):
        d_str = datetime.now().strftime("%Y-%m-%d")
        bk_p = os.path.join(backup_dir, f"lab_data_{d_str}.db")
        if not os.path.exists(bk_p):
            try:
                shutil.copy(db_path, bk_p)
            except:
                pass

def run_query(q, p=(), fetch=False):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute(q, p)
    if fetch:
        d = c.fetchall()
        cols = [desc[0] for desc in c.description]
        conn.close()
        return pd.DataFrame(d, columns=cols)
    conn.commit()
    conn.close()

def insert_task_record(date_str: str, task_name: str, category: str, details: str, tags: str) -> int:
    """插入一条任务记录并返回自增 ID"""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO tasks (date, task_name, category, is_done, details, tags) VALUES (?, ?, ?, ?, ?, ?)",
        (date_str, task_name, category, 0, details or "", tags or "")
    )
    new_id = cur.lastrowid
    conn.commit()
    conn.close()
    return new_id

def get_distinct_tags():
    """获取现有标签下拉选项"""
    df = run_query("SELECT tags FROM tasks WHERE tags IS NOT NULL AND TRIM(tags)!=''", fetch=True)
    if df.empty:
        return []
    tags = set()
    for raw in df["tags"]:
        if not raw:
            continue
        parts = re.split(r"[,，\\s]+", str(raw))
        for part in parts:
            part = part.strip()
            if part:
                tags.add(part)
    return sorted(tags)

# ==================== 优化的历史记录导入 ====================
def import_legacy_records_preserve_original(files, *, default_category: str, default_tags: str, default_date, prefer_filename_date: bool = True, use_ai_metadata: bool = True):
    """
    优化版本：保留原始记录内容，AI只提取元数据
    """
    results = []
    if not files:
        return results
    
    if isinstance(default_date, datetime):
        fallback_date = default_date
    else:
        fallback_date = datetime.combine(default_date, datetime.min.time())
    
    client = get_ai_client() if use_ai_metadata else None
    
    for file_item in files:
        name = getattr(file_item, "name", "legacy_record")
        ext = os.path.splitext(name)[1].lower()
        data = None
        
        try:
            if hasattr(file_item, "getvalue"):
                data = file_item.getvalue()
            elif hasattr(file_item, "read"):
                data = file_item.read()
        except Exception:
            data = None
        
        if not data:
            results.append({"file": name, "success": False, "message": "无法读取文件内容"})
            continue
        
        try:
            # 提取原始文本内容
            if ext in (".md", ".markdown", ".txt", ".csv", ".tsv"):
                original_text = _decode_text_full(data)
                if ext in (".csv", ".tsv"):
                    original_text = f"```\n{original_text}\n```"
            elif ext in (".docx", ".doc", ".rtf"):
                original_text = convert_document_bytes_to_markdown(data, name, ext)
            elif ext in LEGACY_IMAGE_EXTS:
                results.append({"file": name, "success": False, "message": "请将图片嵌入文档一起导入"})
                continue
            else:
                original_text = _decode_text_full(data)
            
            original_text = (original_text or "").strip()
            if not original_text:
                results.append({"file": name, "success": False, "message": "未解析出内容"})
                continue
            
            # 提取元数据
            date_str = guess_record_date_from_filename(name, fallback_date) if prefer_filename_date else fallback_date.strftime("%Y-%m-%d")
            task_name = build_task_name_from_filename(name)
            category = default_category
            tags = default_tags
            
            # 使用AI提取更准确的元数据（可选）
            if client and use_ai_metadata:
                metadata = ai_extract_metadata(client, original_text[:1000])  # 只分析前1000字符
                if metadata:
                    task_name = metadata.get('task_name', task_name)
                    category = metadata.get('category', category)
                    tags = metadata.get('tags', tags)
                    # 如果AI提取了日期，使用它
                    if metadata.get('date'):
                        try:
                            datetime.strptime(metadata['date'], '%Y-%m-%d')
                            date_str = metadata['date']
                        except:
                            pass
            
            # 插入记录，原始内容一字不改
            new_id = insert_task_record(date_str, task_name, category, original_text, tags)
            results.append({
                "file": name, 
                "success": True, 
                "task_id": new_id, 
                "date": date_str,
                "task_name": task_name,
                "category": category,
                "tags": tags,
                "content_preview": original_text[:100] + "..." if len(original_text) > 100 else original_text
            })
            
        except Exception as exc:
            results.append({"file": name, "success": False, "message": str(exc)})
    
    return results

# ==================== 导出功能 ====================
def build_record_markdown(row):
    row = normalize_task_row(row)
    details = row.get("details") or "(暂无实验记录)"
    md = [
        f"# {row.get('task_name', '实验记录')}",
        "",
        f"- 日期：{row.get('date', '-')}",
        f"- 类型：{row.get('category', '-')}",
        f"- 标签：{row.get('tags') or '-'}",
        "",
        "## 实验记录",
        details
    ]
    return "\n".join(md)

def build_record_docx_bytes(row):
    row = normalize_task_row(row)
    doc = Document()
    doc.add_heading(row.get("task_name", "实验记录"), level=1)
    doc.add_paragraph(f"日期：{row.get('date', '-')}")
    doc.add_paragraph(f"类型：{row.get('category', '-')}")
    doc.add_paragraph(f"标签：{row.get('tags') or '-'}")
    doc.add_heading("实验记录", level=2)
    doc.add_paragraph(row.get("details") or "(暂无实验记录)")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_record_exports(row):
    row = normalize_task_row(row)
    base = sanitize_filename(f"{row.get('date', '')}_{row.get('task_name', 'record')}")
    markdown_bytes = build_record_markdown(row).encode("utf-8")
    docx_bytes = build_record_docx_bytes(row)
    return [
        ("MD", f"{base}.md", markdown_bytes, "text/markdown"),
        ("DOCX", f"{base}.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ]

# ==================== Streamlit UI 组件 ====================
@st.dialog("📅 快速添加日程", width="small")
def show_add_task_dialog(default_date_str):
    """快速添加任务对话框"""
    try:
        dd = datetime.strptime(default_date_str, "%Y-%m-%d").date()
    except:
        dd = datetime.now().date()
    
    with st.form("quick_add"):
        st.write(f"日期：**{dd}**")
        col1, col2 = st.columns([3, 1])
        task_name = col1.text_input("内容", placeholder="任务名称")
        category = col2.selectbox("类型", ["科研", "临床", "课程", "其他"])
        tags = st.text_input("标签", "#日常")
        
        if st.form_submit_button("添加", use_container_width=True):
            if task_name.strip():
                run_query(
                    "INSERT INTO tasks (date, task_name, category, is_done, details, tags) VALUES (?,?,?,?,?,?)",
                    (dd, task_name.strip(), category, 0, "", tags.strip())
                )
                st.success("✅ 任务已添加")
                time.sleep(0.5)
                st.rerun()

@st.dialog("📌 任务详情", width="medium")
def show_event_action_dialog(task_id):
    """任务详情对话框"""
    df = run_query("SELECT * FROM tasks WHERE id=?", (task_id,), fetch=True)
    if df.empty:
        st.error("未找到该任务")
        return
    
    row = df.iloc[0]
    st.markdown(f"### {row['task_name']}")
    
    record_status = "✅ 已填写实验记录" if (row['details'] or "").strip() else "🕒 暂未填写实验记录"
    st.caption(f"{record_status} · 标签：{row['tags'] or '-'}")
    
    try:
        date_value = datetime.strptime(str(row['date']), "%Y-%m-%d").date()
    except:
        date_value = datetime.now().date()
    
    with st.form(f"edit_task_{task_id}"):
        name = st.text_input("任务名称", row['task_name'])
        date_input = st.date_input("日期", value=date_value)
        category_options = ["科研", "临床", "课程", "其他"]
        category = st.selectbox("类型", category_options, 
                               index=category_options.index(row['category']) if row['category'] in category_options else 0)
        tags = st.text_input("标签", row['tags'] or "")
        
        submitted = st.form_submit_button("保存修改", use_container_width=True)
        if submitted:
            run_query(
                "UPDATE tasks SET date=?, task_name=?, category=?, tags=? WHERE id=?",
                (date_input.strftime("%Y-%m-%d"), name, category, tags, task_id)
            )
            st.success("✅ 任务已更新")
            time.sleep(0.5)
            st.rerun()
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ 删除任务", type="secondary", use_container_width=True):
            run_query("DELETE FROM tasks WHERE id=?", (task_id,))
            st.success("✅ 已删除")
            time.sleep(0.5)
            st.rerun()
    with col2:
        if st.button("📝 编辑实验记录", type="primary", use_container_width=True):
            st.session_state["open_record_editor_task_id"] = int(task_id)
            st.rerun()

@st.dialog("🧪 编辑实验记录", width="large")
def show_record_editor_dialog(task_id: int):
    """实验记录编辑器"""
    df = run_query("SELECT * FROM tasks WHERE id=?", (task_id,), fetch=True)
    if df.empty:
        st.error("未找到该任务")
        return
    
    row = df.iloc[0]
    st.markdown(f"### [{row['date']}] {row['task_name']}")
    st.caption(f"类型：{row['category'] or '-'} · 当前标签：{row['tags'] or '-'}")
    
    # 初始化状态
    tags_key = f"record_tags_{task_id}"
    details_key = f"record_details_{task_id}"
    init_key = f"record_dialog_init_{task_id}"
    
    if not st.session_state.get(init_key):
        st.session_state[tags_key] = row['tags'] or ""
        st.session_state[details_key] = row['details'] or ""
        st.session_state[init_key] = True
    
    col_main, col_side = st.columns([2, 1])
    
    with col_main:
        st.text_input("标签", key=tags_key)
        st.text_area("实验记录内容", key=details_key, height=350)
        
        # AI润色区域
        st.markdown("#### ✨ AI 润色助手")
        polish_key = f"polish_result_{task_id}"
        feedback_key = f"polish_feedback_{task_id}"
        
        if polish_key not in st.session_state:
            st.session_state[polish_key] = ""
        if feedback_key not in st.session_state:
            st.session_state[feedback_key] = ""
        
        st.text_area("补充要求/反馈", key=feedback_key, height=80)
        
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("✨ 初次润色", key=f"ai_polish_{task_id}"):
                base_text = st.session_state[details_key]
                if not base_text.strip():
                    st.warning("请先填写内容")
                else:
                    client = get_ai_client()
                    if client:
                        extra = (st.session_state[feedback_key] or "").strip() or None
                        with st.spinner("AI 正在润色..."):
                            res = ai_polish_text(client, base_text, extra_instruction=extra)
                        if "Error" not in res:
                            st.session_state[polish_key] = res
                            st.success("✨ 润色完成")
                        else:
                            st.error(res)
        
        with col_btn2:
            disabled = not st.session_state[polish_key]
            if st.button("🪄 根据反馈再润色", key=f"ai_repolish_{task_id}", disabled=disabled):
                client = get_ai_client()
                if client:
                    extra = (st.session_state[feedback_key] or "").strip() or None
                    base_text = st.session_state[polish_key] or st.session_state[details_key]
                    with st.spinner("AI 正在根据反馈润色..."):
                        res = ai_polish_text(client, base_text, extra_instruction=extra)
                    if "Error" not in res:
                        st.session_state[polish_key] = res
                        st.success("✅ 已根据反馈更新")
                    else:
                        st.error(res)
        
        if st.button("💾 保存记录", type="primary", use_container_width=True):
            run_query(
                "UPDATE tasks SET details=?, tags=? WHERE id=?",
                (st.session_state[details_key], st.session_state[tags_key], task_id)
            )
            st.success("✅ 已保存")
            st.session_state[init_key] = False
            time.sleep(0.5)
            st.rerun()
        
        if st.session_state[polish_key]:
            st.text_area("AI 润色结果", st.session_state[polish_key], height=300)
    
    with col_side:
        st.info("📎 附件上传")
        uploads = st.file_uploader("选择文件", accept_multiple_files=True, key=f"upload_{task_id}")
        if uploads:
            for f in uploads:
                save_path, display_name = get_versioned_upload_path(f.name)
                with open(save_path, "wb") as w:
                    w.write(f.getbuffer())
                snippet = f"![{display_name}]({save_path})" if f.type and f.type.startswith("image") else f"[{display_name}]({save_path})"
                st.code(snippet)

@st.dialog("🤖 AI 任务预览与确认", width="large")
def show_ai_confirm_dialog(tasks_data):
    """AI任务确认对话框"""
    st.info("AI 根据您的描述生成了以下计划。请勾选需要导入的项目，也可直接修改内容。")
    
    df = pd.DataFrame(tasks_data)
    if 'task_name' in df.columns:
        df['task_name'] = df['task_name'].apply(lambda x: shorten_task_name(x))
    
    if 'import' not in df.columns:
        df.insert(0, 'import', True)
    
    df['date'] = pd.to_datetime(df['date']).dt.date
    if 'record_outline' not in df.columns:
        df['record_outline'] = ""
    
    edited_df = st.data_editor(
        df,
        column_config={
            "import": st.column_config.CheckboxColumn("导入?", width="small"),
            "date": st.column_config.DateColumn("日期", format="YYYY-MM-DD"),
            "task_name": st.column_config.TextColumn("任务名称"),
            "category": st.column_config.SelectboxColumn("类型", options=["科研", "临床", "课程", "其他"]),
            "tags": st.column_config.TextColumn("标签"),
            "record_outline": st.column_config.TextColumn("实验记录要点", help="将同步写入任务的实验记录详情")
        },
        hide_index=True,
        use_container_width=True,
        num_rows="dynamic"
    )
    
    if st.button("🚀 确认一键导入", type="primary", use_container_width=True):
        count = 0
        for index, row in edited_df.iterrows():
            if row['import']:
                outline = row.get('record_outline', "")
                if pd.isna(outline):
                    outline = ""
                run_query(
                    "INSERT INTO tasks (date, task_name, category, is_done, details, tags) VALUES (?, ?, ?, ?, ?, ?)",
                    (row['date'], row['task_name'], row['category'], 0, outline, row['tags'])
                )
                count += 1
        st.success(f"✅ 成功导入 {count} 条任务！")
        time.sleep(1)
        st.rerun()

# ==================== 主应用 ====================
def setup_page_config():
    """设置页面配置"""
    st.set_page_config(
        page_title="Lab Diary AI - 智能实验记录工具",
        layout="wide",
        initial_sidebar_state="expanded",
        page_icon="🔬"
    )
    
    # 自定义CSS
    st.markdown(f"""
    <style>
    /* 全局样式 */
    .stApp {{
        font-family: {FONTS['family']};
        background-color: {COLORS['background']};
    }}
    
    /* 标题样式 */
    h1, h2, h3 {{
        color: {COLORS['primary']};
        font-weight: 600;
    }}
    
    /* 按钮样式 */
    .stButton > button {{
        border-radius: 8px;
        font-weight: 500;
        transition: all 0.3s ease;
    }}
    
    .stButton > button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }}
    
    /* 卡片样式 */
    .stCard {{
        border-radius: 12px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }}
    
    .stCard:hover {{
        transform: translateY(-4px);
        box-shadow: 0 8px 24px rgba(0,0,0,0.15);
    }}
    
    /* 输入框样式 */
    .stTextInput > div > div > input {{
        border-radius: 8px;
        border: 1px solid #e2e8f0;
    }}
    
    .stTextInput > div > div > input:focus {{
        border-color: {COLORS['accent']};
        box-shadow: 0 0 0 2px rgba(59, 130, 246, 0.2);
    }}
    
    /* 侧边栏样式 */
    [data-testid="stSidebar"] {{
        background-color: white;
        box-shadow: 2px 0 8px rgba(0,0,0,0.1);
    }}
    
    /* 日历样式覆盖 */
    .fc-daygrid-event {{
        height: auto !important;
    }}
    
    .fc-daygrid-event .fc-event-main,
    .fc-daygrid-event .fc-event-title {{
        white-space: normal !important;
        overflow-wrap: anywhere;
        line-height: 1.2;
    }}
    
    /* 响应式设计 */
    @media (max-width: 768px) {{
        [data-testid="stSidebar"] {{
            width: 100% !important;
        }}
    }}
    </style>
    """, unsafe_allow_html=True)

def render_sidebar():
    """渲染侧边栏"""
    with st.sidebar:
        storage = get_storage_paths()
        # Logo和标题
        st.markdown(f"""
        <div style="text-align: center; padding: 20px 0;">
            <h1 style="color: {COLORS['primary']}; font-size: 24px; margin: 0;">🔬 Lab Diary AI</h1>
            <p style="color: {COLORS['secondary']}; font-size: 12px; margin: 5px 0 0 0;">智能实验记录管理</p>
        </div>
        """, unsafe_allow_html=True)
        if storage.get("user_label") and storage["user_label"] != "local":
            st.caption(f"当前用户：{storage['user_label']}")
        
        st.divider()
        
        # 导航菜单
        nav_pages = ["📅 日历总览", "📖 实验记录"]
        page = st.radio("导航", nav_pages, key="nav_page")
        
        st.divider()
        
        # AI助手面板
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, {COLORS['accent']}20, {COLORS['info']}20); 
                    padding: 16px; border-radius: 12px; margin-bottom: 16px;">
            <h3 style="color: {COLORS['primary']}; margin: 0 0 8px 0;">🤖 AI 智能助手</h3>
            <p style="color: {COLORS['secondary']}; font-size: 12px; margin: 0;">
                用自然语言描述您的计划，AI将自动创建任务
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # AI输入
        prompt_key = "ai_schedule_prompt"
        user_prompt = st.text_area(
            "描述您的计划",
            key=prompt_key,
            height=120,
            placeholder="例如：明天开始连续3天测体重，下周五处死取脑"
        )
        
        # 参考文件上传
        uploaded_files = st.file_uploader(
            "参考文件（可选）",
            accept_multiple_files=True,
            key="ai_schedule_files"
        )
        
        # 生成任务按钮
        if st.button("⚡ 生成任务预览", type="primary", use_container_width=True):
            if not user_prompt.strip():
                st.warning("请先输入计划描述")
            else:
                client = get_ai_client()
                if not client:
                    st.error("AI 服务未配置")
                else:
                    with st.spinner("AI 正在分析..."):
                        tasks = ai_parse_schedule(client, user_prompt)
                        if tasks:
                            for task in tasks:
                                task['task_name'] = shorten_task_name(task.get('task_name', ''))
                            show_ai_confirm_dialog(tasks)
                        else:
                            st.error("AI 未能识别出任务，请尝试换个说法")
        
        return page

def render_calendar_page():
    """渲染日历页面"""
    st.markdown(f"""
    <div style="background: white; padding: 24px; border-radius: 16px; box-shadow: 0 4px 12px rgba(0,0,0,0.05);">
        <h1 style="color: {COLORS['primary']}; margin: 0 0 16px 0;">📅 工作日程总览</h1>
    </div>
    """, unsafe_allow_html=True)
    
    # 搜索栏
    col_search, col_filter = st.columns([3, 1])
    search_term = col_search.text_input(
        "🔍 搜索",
        key="calendar_search",
        placeholder="任务名称 / 实验记录 / 标签"
    )
    category_filter = col_filter.selectbox(
        "类别筛选",
        ["全部", "科研", "临床", "课程", "其他"],
        key="calendar_category"
    )
    
    # 日历配置
    cal_ops = {
        "headerToolbar": {
            "left": "today prev,next",
            "center": "title",
            "right": "dayGridMonth,dayGridWeek,dayGridDay"
        },
        "initialView": "dayGridMonth",
        "timeZone": "UTC",
        "buttonText": {"today": "今天", "dayGridMonth": "月", "dayGridWeek": "周", "dayGridDay": "日"},
        "selectable": True,
        "navLinks": False,
        "editable": False,
        "height": 600
    }
    
    # 获取任务数据（支持搜索/筛选）
    where_parts = []
    params: list = []
    if search_term:
        wildcard = f"%{search_term}%"
        where_parts.append("(task_name LIKE ? OR details LIKE ? OR tags LIKE ?)")
        params.extend([wildcard, wildcard, wildcard])
    if category_filter and category_filter != "全部":
        where_parts.append("category=?")
        params.append(category_filter)

    where_sql = (" WHERE " + " AND ".join(where_parts)) if where_parts else ""
    df = run_query("SELECT * FROM tasks" + where_sql + " ORDER BY date", tuple(params), fetch=True)
    df_list = run_query("SELECT * FROM tasks" + where_sql + " ORDER BY date DESC, id DESC", tuple(params), fetch=True)
    events = []
    
    if not df.empty:
        category_color = {
            "科研": COLORS["research"],
            "临床": COLORS["clinical"],
            "课程": COLORS["course"],
            "其他": COLORS["other"],
        }
        for _, r in df.iterrows():
            task_id = int(r['id'])
            color = category_color.get(str(r.get("category", "")).strip(), COLORS["other"])
            
            details_text = (r['details'] or "").strip()
            record_done = bool(details_text)
            prefix = "✅ " if record_done else "⬜ "
            
            event = {
                "id": str(task_id),
                "title": prefix + r['task_name'],
                "start": r['date'],
                "backgroundColor": color,
                "borderColor": color,
                "allDay": True,
                "extendedProps": {
                    "task_id": task_id,
                    "task_name": r['task_name'],
                    "date": r['date'],
                    "category": r['category'],
                    "tags": r['tags'] or "",
                    "is_done": bool(r['is_done']),
                    "details_filled": record_done,
                    "details_preview": details_text[:80] + "..." if len(details_text) > 80 else details_text
                }
            }
            events.append(event)
    
    # 渲染日历（移除右侧快速统计栏）
    cal = calendar(
        events=events,
        options=cal_ops,
        callbacks=['dateClick', 'eventClick', 'eventMouseEnter'],
        key='main_calendar'
    )

    # 处理日历回调
    callback_type = cal.get("callback")
    if callback_type == "dateClick":
        d_str = cal["dateClick"].get("dateStr") or cal["dateClick"].get("date")
        if d_str:
            if "T" in d_str:
                d_str = d_str.split("T")[0]
            show_add_task_dialog(d_str)
    elif callback_type == "eventClick":
        event_payload = cal.get("eventClick", {}).get("event", {})
        props = event_payload.get("extendedProps", {})
        task_id = props.get("task_id") or event_payload.get("id")
        if task_id is not None:
            show_event_action_dialog(int(str(task_id)))

    st.divider()
    st.subheader("📋 任务总列表")

    with st.form("task_quick_add"):
        c1, c2, c3, c4 = st.columns([2, 1, 1, 1])
        new_name = c1.text_input("任务名称", placeholder="例如：测体重/灌胃/行为学测试")
        new_date = c2.date_input("日期", value=datetime.now().date())
        new_category = c3.selectbox("类型", ["科研", "临床", "课程", "其他"])
        new_tags = c4.text_input("标签", value="#日常")
        new_details = st.text_input("备注（可选）")
        submitted = st.form_submit_button("添加任务", use_container_width=True)
        if submitted:
            if not new_name.strip():
                st.warning("请填写任务名称")
            else:
                run_query(
                    "INSERT INTO tasks (date, task_name, category, is_done, details, tags) VALUES (?, ?, ?, ?, ?, ?)",
                    (new_date.strftime("%Y-%m-%d"), new_name.strip(), new_category, 0, new_details.strip(), new_tags.strip())
                )
                st.success("任务已添加")
                st.rerun()

    if df_list.empty:
        st.info("暂无任务数据")
        return

    df_list['date'] = pd.to_datetime(df_list['date']).dt.date
    header_cols = st.columns([0.7, 1.1, 0.8, 1.3, 3, 1.6])
    header_cols[0].markdown("**完成**")
    header_cols[1].markdown("**日期**")
    header_cols[2].markdown("**类型**")
    header_cols[3].markdown("**标签/记录**")
    header_cols[4].markdown("**内容**")
    header_cols[5].markdown("**操作**")

    for _, row in df_list.iterrows():
        with st.container():
            c1, c2, c3, c4, c5, c6 = st.columns([0.7, 1.1, 0.8, 1.3, 3, 1.6])
            done_val = bool(row['is_done'])
            checked = c1.checkbox("", done_val, key=f"task_done_{row['id']}")
            if checked != done_val:
                run_query("UPDATE tasks SET is_done=? WHERE id=?", (1 if checked else 0, row['id']))
                st.rerun()

            c2.text(str(row['date']))
            c3.caption(row['category'])
            record_flag = "✅ 已写" if (row['details'] or "").strip() else "✏️ 待写"
            c4.markdown(f"{row['tags'] or '-'} · {record_flag}")
            if row['is_done']:
                c5.markdown(f"~~{row['task_name']}~~")
            else:
                c5.text(row['task_name'])

            action_col, record_col = c6.columns(2)
            if action_col.button("详情", key=f"task_detail_{row['id']}"):
                show_event_action_dialog(int(row['id']))
            if record_col.button("记录", key=f"task_record_{row['id']}"):
                show_record_editor_dialog(int(row['id']))

            st.markdown("<hr style='margin:0.2em 0;opacity:0.1'>", unsafe_allow_html=True)

# ==================== 主函数 ====================
def main():
    """主函数"""
    setup_page_config()
    init_and_migrate_db()
    auto_backup()
    
    # 初始化会话状态
    if "nav_page" not in st.session_state:
        st.session_state["nav_page"] = "📅 日历总览"
    
    # 渲染侧边栏并获取当前页面
    page = render_sidebar()

    # 避免 Dialog 嵌套：在对话框内只设置标记并 rerun，真正打开编辑器在主渲染阶段完成
    pending_record_task_id = st.session_state.pop("open_record_editor_task_id", None)
    if pending_record_task_id is not None:
        show_record_editor_dialog(int(pending_record_task_id))
    
    # 根据页面渲染内容
    if page == "📅 日历总览":
        render_calendar_page()
    elif page == "📖 实验记录":
        render_archive_page()

if __name__ == "__main__":
    main()
# 继续添加缺失的函数

# ==================== 辅助函数（继续） ====================
def _safe_date_from_parts(year: str, month: str, day: str) -> str | None:
    """从年月日组件安全创建日期字符串"""
    try:
        dt = datetime(year=int(year), month=int(month), day=int(day))
        return dt.strftime("%Y-%m-%d")
    except Exception:
        return None

def guess_record_date_from_filename(filename: str, fallback_date: datetime) -> str:
    """根据文件名中的日期信息推测日志日期"""
    base = os.path.basename(filename)
    stem = os.path.splitext(base)[0]
    patterns = [
        r"(20\d{2})[-_/\.](\d{1,2})[-_/\.](\d{1,2})",
        r"(20\d{2})(\d{2})(\d{2})",
        r"(20\d{2})年(\d{1,2})月(\d{1,2})日",
    ]
    for pattern in patterns:
        match = re.search(pattern, stem)
        if match:
            guess = _safe_date_from_parts(*match.groups())
            if guess:
                return guess
    return fallback_date.strftime("%Y-%m-%d")

def build_task_name_from_filename(filename: str) -> str:
    """将文件名转为易读的任务标题"""
    base = os.path.basename(filename)
    stem = os.path.splitext(base)[0]
    stem = re.sub(r"(20\d{2}[^\d]?\d{1,2}[^\d]?\d{1,2})", " ", stem)
    stem = re.sub(r"[_\-]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    if not stem:
        stem = "历史记录"
    return shorten_task_name(stem)

# ==================== DOCX 处理函数 ====================
def _iter_docx_block_items(parent):
    """迭代DOCX文档中的块元素"""
    if isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        parent_element = parent.element.body
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def _paragraph_is_list(paragraph: Paragraph) -> bool:
    """判断段落是否为列表项"""
    p = paragraph._p
    if p is None or p.pPr is None:
        return False
    return p.pPr.numPr is not None

def _paragraph_is_code(paragraph: Paragraph) -> bool:
    """判断段落是否为代码块"""
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    code_keywords = ("code", "等宽", "monospace")
    if any(token in style_name for token in code_keywords):
        return True
    for run in paragraph.runs:
        font = getattr(run, "font", None)
        if font and font.name:
            fname = font.name.lower()
            if any(token in fname for token in ("consolas", "courier", "monospace")):
                return True
    return False

def _heading_level_from_style(style_name: str) -> int:
    """从样式名称提取标题级别"""
    match = re.search(r"(\d+)", style_name)
    if match:
        try:
            return max(1, min(6, int(match.group(1))))
        except ValueError:
            pass
    return 1

def _docx_paragraph_to_markdown(paragraph: Paragraph) -> str:
    """将DOCX段落转换为Markdown"""
    text = paragraph.text.strip()
    if not text:
        return ""
    style_name = (paragraph.style.name or "").lower() if paragraph.style and paragraph.style.name else ""
    if "heading" in style_name or "标题" in style_name:
        level = _heading_level_from_style(style_name)
        return f"{'#' * level} {text}"
    if _paragraph_is_code(paragraph):
        return f"```\n{text}\n```"
    if _paragraph_is_list(paragraph) or "list" in style_name or "列表" in style_name:
        return f"- {text}"
    return text

def _docx_table_to_markdown(table: Table) -> str:
    """将DOCX表格转换为Markdown"""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            cell_text = cell_text.replace("\n", "<br>")
            cells.append(cell_text or " ")
        rows.append(cells)
    if not rows:
        return ""
    header = rows[0]
    divider = ["---"] * len(header)
    body = rows[1:] or [[]]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(divider) + " |",
    ]
    if body and body[0]:
        for row in body:
            padded = row + [" "] * (len(header) - len(row))
            lines.append("| " + " | ".join(padded[:len(header)]) + " |")
    return "\n".join(lines)

def _collect_docx_image_markdown(docx_bytes: bytes, origin_name: str) -> list[str]:
    """从DOCX中提取图片"""
    images = []
    label_prefix = sanitize_filename(os.path.splitext(origin_name)[0] or "legacy_doc")
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        for entry in archive.infolist():
            if entry.is_dir():
                continue
            if not entry.filename.startswith("word/media/"):
                continue
            data = archive.read(entry.filename)
            img_name = os.path.basename(entry.filename)
            images.append(_persist_image_as_markdown(data, f"{label_prefix}_{img_name}"))
    return images

def _pandoc_convert(data_bytes: bytes, source_ext: str, target: str) -> bytes | None:
    """使用Pandoc转换文档"""
    if not HAS_PYPANDOC:
        return None
    suffix = source_ext if source_ext.startswith(".") else f".{source_ext}"
    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp_in.write(data_bytes)
    tmp_in.close()
    try:
        if target == "docx":
            tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tmp_out_path = tmp_out.name
            tmp_out.close()
            pypandoc.convert_file(tmp_in.name, to="docx", format=source_ext.lstrip("."), outputfile=tmp_out_path)
            with open(tmp_out_path, "rb") as fh:
                return fh.read()
        else:
            result = pypandoc.convert_file(tmp_in.name, to=target, format=source_ext.lstrip("."))
            return result.encode("utf-8")
    except (OSError, RuntimeError):
        return None
    finally:
        try:
            os.remove(tmp_in.name)
        except OSError:
            pass

def _convert_doc_via_win32(data_bytes: bytes) -> bytes | None:
    """在Windows环境下将DOC转为DOCX"""
    if not HAS_WIN32_COM:
        return None
    temp_dir = tempfile.mkdtemp()
    doc_path = os.path.join(temp_dir, "legacy.doc")
    docx_path = os.path.join(temp_dir, "legacy.docx")
    with open(doc_path, "wb") as fh:
        fh.write(data_bytes)
    converted = None
    word = None
    doc_obj = None
    initialized = False
    try:
        pythoncom.CoInitialize()
        initialized = True
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_obj = word.Documents.Open(doc_path)
        doc_obj.SaveAs(docx_path, FileFormat=16)
        doc_obj.Close(False)
        with open(docx_path, "rb") as fh:
            converted = fh.read()
    except Exception:
        converted = None
    finally:
        if doc_obj is not None:
            try:
                doc_obj.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if initialized:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
        for path in (doc_path, docx_path):
            if os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass
        try:
            os.rmdir(temp_dir)
        except OSError:
            pass
    return converted

# ==================== 归档页面 ====================
def render_archive_page():
    """渲染实验记录页面"""
    st.markdown(f"""
    <div style="background: white; padding: 24px; border-radius: 16px; box-shadow: 0 4px 12px rgba(0,0,0,0.05);">
        <h1 style="color: {COLORS['primary']}; margin: 0 0 16px 0;">📖 实验记录</h1>
        <p style="color: {COLORS['secondary']}; margin: 0;">管理和导出您的实验记录</p>
    </div>
    """, unsafe_allow_html=True)
    
    # 搜索和筛选
    col_search, col_tag = st.columns([3, 1])
    search_term = col_search.text_input("🔍 搜索", key="archive_search")
    
    tags_available = ["全部"] + get_distinct_tags()
    tag_choice = col_tag.selectbox("标签筛选", tags_available, key="archive_tag")
    
    # 一键迁移历史记录
    with st.expander("🪄 一键迁移历史记录", expanded=False):
        st.caption("支持 Markdown / Word / TXT / CSV 等格式，保留原始记录内容")
        
        legacy_files = st.file_uploader(
            "选择旧实验记录文件",
            accept_multiple_files=True,
            type=["md", "markdown", "txt", "csv", "tsv", "doc", "docx", "rtf"],
            key="legacy_import_files"
        )
        
        col1, col2, col3 = st.columns(3)
        legacy_category = col1.selectbox("导入类别", ["科研", "临床", "课程", "其他"], key="legacy_category")
        legacy_tags = col2.text_input("统一标签", "#历史记录", key="legacy_tags")
        legacy_date = col3.date_input("默认日期", datetime.now().date(), key="legacy_date")
        
        use_ai = st.checkbox("使用AI提取元数据（推荐）", value=True, key="use_ai_metadata")
        filename_date = st.checkbox("尝试根据文件名推断日期", value=True, key="filename_date")
        
        if st.button("🚀 开始迁移", type="primary", use_container_width=True):
            if not legacy_files:
                st.warning("请先选择至少一个文件")
            else:
                with st.spinner("正在解析并导入历史记录..."):
                    import_results = import_legacy_records_preserve_original(
                        legacy_files,
                        default_category=legacy_category,
                        default_tags=legacy_tags,
                        default_date=legacy_date,
                        prefer_filename_date=filename_date,
                        use_ai_metadata=use_ai
                    )
                
                # 显示结果
                success_items = [item for item in import_results if item.get("success")]
                failure_items = [item for item in import_results if not item.get("success")]
                
                if success_items:
                    st.success(f"✅ 成功导入 {len(success_items)} 条记录")
                    for item in success_items:
                        with st.expander(f"✅ {item['file']}"):
                            st.write(f"**任务名**: {item['task_name']}")
                            st.write(f"**日期**: {item['date']}")
                            st.write(f"**类别**: {item['category']}")
                            st.write(f"**标签**: {item['tags']}")
                            st.write(f"**预览**: {item['content_preview']}")
                
                if failure_items:
                    st.error(f"❌ {len(failure_items)} 个文件导入失败")
                    for item in failure_items:
                        st.write(f"⚠️ {item['file']}: {item.get('message', '未知错误')}")
    
    # 自动生成周报
    with st.expander("🗓️ 自动生成周报", expanded=False):
        reference_date = st.date_input("参考日期", datetime.now().date(), key="weekly_date")
        
        if st.button("📄 生成周报", key="weekly_btn"):
            start_date = (reference_date - timedelta(days=6)).strftime("%Y-%m-%d")
            end_date = reference_date.strftime("%Y-%m-%d")
            
            report_df = run_query(
                "SELECT date, task_name, details, tags FROM tasks WHERE category='科研' AND details!='' AND date BETWEEN ? AND ? ORDER BY date",
                (start_date, end_date),
                fetch=True
            )
            
            if report_df.empty:
                st.warning("所选时间段内暂无实验记录")
            else:
                records = report_df.to_dict('records')
                client = get_ai_client()
                
                if client:
                    with st.spinner("AI 正在整理周报..."):
                        report_text = ai_generate_weekly_report(client, records, start_date, end_date)
                else:
                    report_text = build_weekly_report_fallback(records, start_date, end_date)
                
                st.text_area("周报内容", report_text, height=300)
                st.download_button(
                    "📥 导出周报",
                    report_text.encode("utf-8"),
                    file_name=f"weekly_report_{end_date}.md",
                    mime="text/markdown",
                    use_container_width=True
                )
    
    # 查询记录
    base_sql = "SELECT * FROM tasks WHERE category='科研' AND details!=''"
    params = []
    
    if search_term:
        wildcard = f"%{search_term}%"
        base_sql += " AND (task_name LIKE ? OR details LIKE ? OR tags LIKE ?)"
        params.extend([wildcard, wildcard, wildcard])
    
    if tag_choice and tag_choice != "全部":
        base_sql += " AND tags LIKE ?"
        params.append(f"%{tag_choice}%")
    
    df = run_query(base_sql + " ORDER BY date DESC", tuple(params), fetch=True)
    
    if df.empty:
        st.info("📭 暂时没有符合条件的实验记录")
    else:
        # 批量导出
        st.markdown("### 📤 批量导出")
        archive_exports = get_archive_exports(df.to_dict('records'))
        
        if archive_exports:
            exp_cols = st.columns(len(archive_exports))
            for col, item in zip(exp_cols, archive_exports):
                label, fname, data, mime = item
                with col:
                    st.download_button(
                        f"📄 导出{label}",
                        data,
                        file_name=fname,
                        mime=mime,
                        use_container_width=True
                    )
        
        st.divider()
        
        # 记录列表
        st.markdown(f"### 📋 实验记录列表 ({len(df)}条)")
        
        for _, r in df.iterrows():
            with st.expander(f"**{r['date']}** | {r['task_name']}", expanded=False):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.caption(f"🏷️ 标签：{r['tags'] or '-'} · 📂 类型：{r['category']}")
                    st.markdown(r['details'])
                with col2:
                    if st.button("📝 编辑", key=f"edit_{r['id']}", use_container_width=True):
                        show_record_editor_dialog(int(r['id']))
                    
                    # 单个导出
                    exports = get_record_exports(r)
                    for label, fname, data, mime in exports:
                        st.download_button(
                            f"📄 {label}",
                            data,
                            file_name=fname,
                            mime=mime,
                            key=f"export_{r['id']}_{label}",
                            use_container_width=True
                        )

def render_analytics_page():
    """数据分析页面（已下线）"""
    st.info("“数据分析”功能已下线（不再维护）。")
    return
    st.markdown(f"""
    <div style="background: white; padding: 24px; border-radius: 16px; box-shadow: 0 4px 12px rgba(0,0,0,0.05);">
        <h1 style="color: {COLORS['primary']}; margin: 0 0 16px 0;">📊 数据分析</h1>
        <p style="color: {COLORS['secondary']}; margin: 0;">可视化您的科研工作数据</p>
    </div>
    """, unsafe_allow_html=True)
    
    # 获取数据
    df = run_query("SELECT * FROM tasks ORDER BY date", fetch=True)
    
    if df.empty:
        st.info("📭 暂无数据可供分析")
        return
    
    # 统计卡片
    st.markdown("### 📈 总体统计")
    col1, col2, col3, col4 = st.columns(4)
    
    total_tasks = len(df)
    completed_tasks = len(df[df['is_done'] == 1])
    research_tasks = len(df[df['category'] == '科研'])
    this_week_tasks = len(df[df['date'] >= (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')])
    
    with col1:
        st.metric("总任务数", total_tasks)
    with col2:
        st.metric("已完成", completed_tasks)
    with col3:
        st.metric("科研任务", research_tasks)
    with col4:
        st.metric("本周任务", this_week_tasks)
    
    # 图表区域
    col_chart1, col_chart2 = st.columns(2)
    
    with col_chart1:
        st.markdown("### 📅 工作量趋势")
        workload_chart = create_workload_chart(df)
        if workload_chart:
            st.plotly_chart(workload_chart, use_container_width=True)
    
    with col_chart2:
        st.markdown("### 🥧 类别分布")
        category_chart = create_category_pie_chart(df)
        if category_chart:
            st.plotly_chart(category_chart, use_container_width=True)
    
    # 标签云
    st.markdown("### 🏷️ 标签分析")
    all_tags = []
    for tags in df['tags'].dropna():
        tag_list = re.split(r'[,，\s]+', str(tags))
        all_tags.extend([tag.strip() for tag in tag_list if tag.strip() and tag.strip().startswith('#')])
    
    if all_tags:
        tag_counts = Counter(all_tags)
        top_tags = dict(tag_counts.most_common(20))
        
        # 创建标签云可视化
        tag_df = pd.DataFrame(list(top_tags.items()), columns=['标签', '次数'])
        fig = px.bar(tag_df, x='次数', y='标签', orientation='h', 
                     title='常用标签统计 (Top 20)')
        fig.update_layout(
            font=dict(family=FONTS['family'], size=12),
            plot_bgcolor='white',
            paper_bgcolor='white'
        )
        st.plotly_chart(fig, use_container_width=True)
    
    # 导出统计数据
    st.markdown("### 📤 导出数据")
    if st.button("📊 导出CSV报告", use_container_width=True):
        csv_data = df.to_csv(index=False).encode('utf-8')
        st.download_button(
            "💾 下载CSV",
            csv_data,
            file_name=f"lab_report_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True
        )

# ==================== 导出功能（继续） ====================
def build_archive_markdown(rows):
    """构建归档Markdown"""
    sections = [build_record_markdown(r) for r in rows]
    return "\n\n---\n\n".join(sections)

def build_archive_docx_bytes(rows):
    """构建归档DOCX"""
    doc = Document()
    for idx, row in enumerate(rows):
        if idx > 0:
            doc.add_page_break()
        row = normalize_task_row(row)
        doc.add_heading(row.get("task_name", "实验记录"), level=1)
        doc.add_paragraph(f"日期：{row.get('date', '-')}")
        doc.add_paragraph(f"类型：{row.get('category', '-')}")
        doc.add_paragraph(f"标签：{row.get('tags') or '-'}")
        doc.add_heading("实验记录", level=2)
        doc.add_paragraph(row.get("details") or "(暂无实验记录)")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_archive_exports(rows):
    """获取归档导出数据"""
    if not rows:
        return []
    rows = [normalize_task_row(r) for r in rows]
    timestamp = datetime.now().strftime("%Y%m%d")
    base = f"lab_archive_{timestamp}"
    md_bytes = build_archive_markdown(rows).encode("utf-8")
    docx_bytes = build_archive_docx_bytes(rows)
    return [
        ("MD", f"{base}.md", md_bytes, "text/markdown"),
        ("DOCX", f"{base}.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    ]

def build_weekly_report_fallback(records, start_date, end_date):
    """无AI时的简单周报拼接"""
    lines = [f"# 周报（{start_date} ~ {end_date}）", ""]
    for row in records:
        snippet = (row.get("details") or "").replace("\n", " ")
        snippet = re.sub(r"\s+", " ", snippet)
        if len(snippet) > 160:
            snippet = snippet[:160] + "…"
        lines.append(f"- {row.get('date', '')} {row.get('task_name', '')}：{snippet}")
    return "\n".join(lines)
